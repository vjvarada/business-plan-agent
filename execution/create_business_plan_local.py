#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Create business plan as LOCAL WORD DOCUMENT (.docx).
Compatible with Google Docs when uploaded.

Usage:
    python create_business_plan_local.py --company "RapidTools" --industry "Manufacturing" --description "Design automation for tooling"
"""

import os
import sys
import argparse
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_business_plan_local(company_name, industry, description):
    """Create a local Word document business plan."""
    doc = Document()
    
    # Title
    title = doc.add_heading(f'{company_name} Business Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Executive Summary
    doc.add_heading('Executive Summary', 1)
    doc.add_paragraph(f'{company_name} operates in the {industry} industry.')
    doc.add_paragraph(description)
    doc.add_paragraph('')
    doc.add_paragraph('This business plan outlines our market opportunity, competitive positioning, and financial projections for the next 5-10 years.')
    
    # Company Overview
    doc.add_heading('Company Overview', 1)
    doc.add_paragraph(f'Company Name: {company_name}')
    doc.add_paragraph(f'Industry: {industry}')
    doc.add_paragraph(f'Mission: {description}')
    
    # Market Analysis
    doc.add_heading('Market Analysis', 1)
    doc.add_paragraph('TAM (Total Addressable Market): TBD')
    doc.add_paragraph('SAM (Serviceable Available Market): TBD')
    doc.add_paragraph('SOM (Serviceable Obtainable Market): TBD')
    doc.add_paragraph('')
    doc.add_paragraph('[Populate with market research using serp_market_research.py]')
    
    # SWOT Analysis
    doc.add_heading('SWOT Analysis', 1)
    
    doc.add_heading('Strengths', 2)
    doc.add_paragraph('[To be populated]')
    
    doc.add_heading('Weaknesses', 2)
    doc.add_paragraph('[To be populated]')
    
    doc.add_heading('Opportunities', 2)
    doc.add_paragraph('[To be populated]')
    
    doc.add_heading('Threats', 2)
    doc.add_paragraph('[To be populated]')
    
    # Business Model
    doc.add_heading('Business Model', 1)
    doc.add_paragraph('Revenue Streams: [To be defined]')
    doc.add_paragraph('Key Partners: [To be defined]')
    doc.add_paragraph('Key Activities: [To be defined]')
    doc.add_paragraph('Customer Segments: [To be defined]')
    
    # Financial Projections
    doc.add_heading('Financial Projections', 1)
    doc.add_paragraph('See accompanying financial model for detailed 10-year projections.')
    doc.add_paragraph('[Link to Excel model will be added after upload]')
    
    # Sources & References
    doc.add_heading('Sources & References', 1)
    doc.add_paragraph('[Market research sources will be listed here]')
    
    # Save
    os.makedirs('.tmp', exist_ok=True)
    filename = f'{company_name.replace(" ", "_")}_business_plan.docx'
    filepath = os.path.join('.tmp', filename)
    doc.save(filepath)
    
    print(f'\n Business plan created!')
    print(f'   File: {filepath}')
    print(f'   Sections: Executive Summary, Market Analysis, SWOT, Business Model, Financial Projections')
    print(f'\nNext steps:')
    print(f'  1. Open in Word to review and edit')
    print(f'  2. Populate with market research data')
    print(f'  3. Upload to Google Drive when ready (use sync_to_cloud.py)')
    
    return filepath

def main():
    parser = argparse.ArgumentParser(description='Create local Word business plan')
    parser.add_argument('--company', required=True, help='Company name')
    parser.add_argument('--industry', required=True, help='Industry sector')
    parser.add_argument('--description', required=True, help='Business description')
    
    args = parser.parse_args()
    
    try:
        filepath = create_business_plan_local(
            company_name=args.company,
            industry=args.industry,
            description=args.description
        )
        return filepath
    except Exception as e:
        print(f'\n Error: {e}')
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == '__main__':
    main()